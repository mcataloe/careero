from __future__ import annotations

import hashlib
import re
import uuid
from dataclasses import dataclass
from datetime import datetime, timezone
from io import BytesIO
from textwrap import wrap
from typing import Literal

from sqlalchemy.orm import Session

from app.models import GeneratedArtifact, User
from app.services.current_user import CurrentUserResolutionError, resolve_current_user
from app.services.integration_adapters import LocalIntegrationAdapter


ArtifactExportFormat = Literal["md", "docx", "pdf"]


class ArtifactExportError(Exception):
    pass


class ArtifactExportNotFoundError(ArtifactExportError):
    pass


class ArtifactExportSeedMissingError(ArtifactExportError):
    pass


class ArtifactExportDependencyError(ArtifactExportError):
    pass


class ArtifactExportValidationError(ArtifactExportError):
    pass


@dataclass(frozen=True)
class ArtifactExportResult:
    content: bytes
    media_type: str
    file_name: str
    content_hash: str


class ArtifactExportService:
    def __init__(self, db: Session) -> None:
        self.db = db
        self.adapter = LocalIntegrationAdapter()

    def get_default_user(self) -> User:
        try:
            return resolve_current_user(self.db)
        except CurrentUserResolutionError as exc:
            raise ArtifactExportSeedMissingError(
                "Default local user is missing; run python -m app.seed"
            ) from exc

    def export_artifact(
        self,
        *,
        artifact_id: uuid.UUID,
        export_format: ArtifactExportFormat,
    ) -> ArtifactExportResult:
        user = self.get_default_user()
        artifact = self.db.get(GeneratedArtifact, artifact_id)
        if (
            artifact is None
            or artifact.user_id != user.id
            or artifact.deleted_at is not None
        ):
            raise ArtifactExportNotFoundError("Artifact not found")
        if export_format not in {"md", "docx", "pdf"}:
            raise ArtifactExportValidationError(
                f"Unsupported artifact export format: {export_format}"
            )

        draft = self.adapter.build_draft(
            operation="export",
            target_type="artifact",
            target_id=str(artifact.id),
            payload={"format": export_format, "artifact_type": artifact.artifact_type},
            provenance={"source": "local_artifact_export"},
        )
        content = _render_export_bytes(
            title=artifact.title,
            content=artifact.content,
            export_format=export_format,
        )
        content_hash = f"sha256:{hashlib.sha256(content).hexdigest()}"
        file_name = _export_file_name(artifact.title, export_format)
        result = ArtifactExportResult(
            content=content,
            media_type=_media_type(export_format),
            file_name=file_name,
            content_hash=content_hash,
        )
        self._record_export(
            artifact=artifact,
            export_format=export_format,
            file_name=file_name,
            content_hash=content_hash,
            adapter_name=draft.adapter_name,
        )
        self.db.commit()
        return result

    def _record_export(
        self,
        *,
        artifact: GeneratedArtifact,
        export_format: ArtifactExportFormat,
        file_name: str,
        content_hash: str,
        adapter_name: str,
    ) -> None:
        exported_at = datetime.now(timezone.utc).isoformat()
        metadata = dict(artifact.artifact_metadata or {})
        contract = metadata.get("contract")
        export_record = {
            "format": export_format,
            "exportedAt": exported_at,
            "fileName": file_name,
            "contentHash": content_hash,
        }

        if isinstance(contract, dict):
            contract = dict(contract)
            export_metadata = list(contract.get("exportMetadata") or [])
            export_metadata.append(export_record)
            contract["exportMetadata"] = export_metadata
            format_metadata = contract.get("formatMetadata")
            if isinstance(format_metadata, dict):
                available_formats = list(format_metadata.get("availableFormats") or [])
                if export_format not in available_formats:
                    available_formats.append(export_format)
                contract["formatMetadata"] = {
                    **format_metadata,
                    "availableFormats": available_formats,
                }
            metadata["contract"] = contract

        export_history = list(metadata.get("export_history") or [])
        export_history.append(
            {
                **export_record,
                "adapterName": adapter_name,
                "localOnly": True,
                "externalMutation": False,
            }
        )
        metadata["export_history"] = export_history
        artifact.artifact_metadata = metadata


def _render_export_bytes(
    *,
    title: str,
    content: str,
    export_format: ArtifactExportFormat,
) -> bytes:
    if export_format == "md":
        return content.encode("utf-8")
    if export_format == "docx":
        return _render_docx(title=title, content=content)
    if export_format == "pdf":
        return _render_pdf(title=title, content=content)
    raise ArtifactExportValidationError(
        f"Unsupported artifact export format: {export_format}"
    )


def _render_docx(*, title: str, content: str) -> bytes:
    try:
        from docx import Document
    except ImportError as exc:
        raise ArtifactExportDependencyError(
            "DOCX export requires the python-docx package"
        ) from exc

    document = Document()
    document.add_heading(title, level=1)
    for line in content.splitlines():
        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
        elif stripped.startswith("#"):
            level = min(len(stripped) - len(stripped.lstrip("#")), 3)
            document.add_heading(stripped.lstrip("#").strip(), level=level)
        elif stripped.startswith(("- ", "* ")):
            document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _render_pdf(*, title: str, content: str) -> bytes:
    lines = [title, ""]
    for source_line in content.splitlines():
        stripped = source_line.strip()
        if not stripped:
            lines.append("")
            continue
        if stripped.startswith("#"):
            stripped = stripped.lstrip("#").strip()
        if stripped.startswith(("- ", "* ")):
            stripped = f"- {stripped[2:].strip()}"
        lines.extend(wrap(stripped, width=92) or [""])
    visible_lines = lines[:52]
    stream_lines = ["BT", "/F1 11 Tf", "50 770 Td", "14 TL"]
    for line in visible_lines:
        stream_lines.append(f"({_escape_pdf_text(line)}) Tj")
        stream_lines.append("T*")
    stream_lines.append("ET")
    stream = "\n".join(stream_lines).encode("latin-1", errors="replace")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(stream)).encode("ascii") + b" >>\nstream\n" + stream + b"\nendstream",
    ]
    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, body in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{index} 0 obj\n".encode("ascii"))
        pdf.extend(body)
        pdf.extend(b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    pdf.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    pdf.extend(
        f"trailer << /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode("ascii")
    )
    return bytes(pdf)


def _escape_pdf_text(value: str) -> str:
    return value.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def _media_type(export_format: ArtifactExportFormat) -> str:
    return {
        "md": "text/markdown; charset=utf-8",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "pdf": "application/pdf",
    }[export_format]


def _export_file_name(title: str, export_format: ArtifactExportFormat) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", title).strip("-").lower()
    if not slug:
        slug = "careero-artifact"
    return f"{slug}.{export_format}"
